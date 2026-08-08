import docx
import sys

sys.stdout.reconfigure(encoding='utf-8')
doc = docx.Document("../Ugwulebo_Lesile_Ngozi_MIT_AWS_GCP_Final_Thesis.docx")

for idx in range(2018, 2258):
    if idx < len(doc.paragraphs):
        para = doc.paragraphs[idx]
        text = para.text.strip()
        if text.startswith("4.8") or text.startswith("File:") or "resource " in text or "name " in text or text == "}":
            print(f"P{idx}: {repr(para.text)}")
