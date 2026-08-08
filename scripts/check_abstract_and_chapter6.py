import docx

doc_path = "../Ugwulebo_Lesile_Ngozi_MIT_AWS_GCP_Final_Thesis.docx"
doc = docx.Document(doc_path)

print("--- ABSTRACT PARAGRAPHS (P23 - P45) ---")
for idx in range(23, 46):
    print(f"P{idx}: {doc.paragraphs[idx].text}")

print("\n--- CHAPTER 6 SYNTHESIS PARAGRAPHS (P5030 - P5080) ---")
for idx in range(5030, min(5080, len(doc.paragraphs))):
    print(f"P{idx}: {doc.paragraphs[idx].text}")
