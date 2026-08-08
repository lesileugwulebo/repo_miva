import docx

doc_path = "../Ugwulebo_Lesile_Ngozi_MIT_AWS_GCP_FINAL_SUBMISSION_READY.docx"
doc = docx.Document(doc_path)

print(f"Deep scanning {len(doc.paragraphs)} paragraphs in {doc_path}...")

contradiction_keywords = [
    "pending", "outstanding", "awaiting", "not yet", "has not been", "must be inserted",
    "substantially achieved", "partially confirmed", "deliberately uses", "result placeholders",
    "to be completed", "after execution", "when completed", "final results should"
]

found_contradictions = []
for idx, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    text_lower = text.lower()
    
    # Check for contradictions
    for kw in contradiction_keywords:
        if kw in text_lower:
            # Exclude legitimate future work / recommendations in section 6.7/6.8
            if not (idx > 5000 and any(k in text_lower for k in ["recommend", "future", "could", "expansion"])):
                found_contradictions.append((idx, kw, text))
                break

print(f"\n--- Contradiction & Pending Language Findings ({len(found_contradictions)}) ---")
for idx, kw, text in found_contradictions[:30]:
    print(f"P{idx} [kw: '{kw}']: {text[:150]}")

# Check duplicate headings
print("\n--- Duplicate Headings Check ---")
headings = {}
duplicate_headings = []
for idx, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    if para.style.name.startswith("Heading") or text.isupper():
        if len(text) > 3 and text in headings:
            duplicate_headings.append((idx, text, headings[text]))
        else:
            headings[text] = idx

print(f"Duplicate Headings Found: {len(duplicate_headings)}")
for idx, text, prev_idx in duplicate_headings:
    print(f"P{idx} duplicates P{prev_idx}: '{text}'")

# Check cloud terminology mix-ups
print("\n--- Cloud Terminology Audit Findings ---")
term_issues = []
for idx, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    if "gcp security group" in text.lower() or "aws nsg" in text.lower() or "microsoft aws" in text.lower() or "expressroute" in text.lower():
        term_issues.append((idx, text))

print(f"Terminology Issues Found: {len(term_issues)}")
for idx, text in term_issues:
    print(f"P{idx}: {text[:150]}")
