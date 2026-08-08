import docx
import re
import sys

sys.stdout.reconfigure(encoding='utf-8')

doc = docx.Document("../Ugwulebo_Lesile_Ngozi_MIT_AWS_GCP_Final_Thesis.docx")

# Pattern for placeholders
placeholder_pattern = re.compile(r"\[[A-Za-z0-9_/ -]{3,}\]")

# Keywords indicating mixed terminology or formatting errors
checks = {
    "azure_prefix": ["awsrm_", "azurerm_", "awsad_"],
    "azure_terminology": ["resource group", "Log Analytics", "Virtual WAN", "system-assigned managed identity", "Network VPC Firewall Rules"],
    "citations": ["Verdet et al. (2025)", "Verdet (2025)"],
    "ligatures": [r"\bdi er\b", r"\bo icial\b", r"\btra ic\b", r"\bo boarding\b", r"\bdi erent\b", r"\bdi erences\b"],
    "disclaimers": ["does not invent", "pending actual execution", "numerical result fields", "remain placeholders"],
    "placeholders": ["screenshot placeholder", "insert screenshot showing"]
}

print("Running refined validation checks on the finalized thesis...")
anomalies_found = 0

for i, para in enumerate(doc.paragraphs):
    text = para.text
    clean_text = " ".join(text.split())
    
    # 1. Check placeholders [INSERT]
    matches = placeholder_pattern.findall(text)
    if matches:
        # Ignore NITDA or allowed guidelines text
        if not any(k in text for k in ["[NITDA]", "[Unit]", "[Service]", "[Install]"]):
            if "guidelines" not in text.lower() and "every field marked" not in text.lower() and "where actual cloud" not in text.lower():
                print(f"[Placeholder] P{i}: matches={matches} -> {repr(clean_text[:120])}")
                anomalies_found += 1

    # 2. Check keywords
    for cat, terms in checks.items():
        for term in terms:
            if cat == "ligatures":
                if re.search(term, clean_text, re.IGNORECASE):
                    print(f"[{cat.upper()}] P{i}: matched pattern {term} -> {repr(clean_text[:120])}")
                    anomalies_found += 1
            else:
                if term.lower() in clean_text.lower():
                    # Double check if it's AWS resource groups or GCP subnets context
                    if "resource group" in term.lower() and ("resourcegroups_group" in clean_text or "aws_resourcegroups_group" in clean_text):
                        continue # allow official AWS Resource Groups resource name
                    if "network vpc firewall rules" in term.lower() and "vpc firewall rules" in clean_text.lower() and "gcp" in clean_text.lower():
                        continue # allow GCP VPC firewall rules
                    print(f"[{cat.upper()}] P{i}: matched '{term}' -> {repr(clean_text[:120])}")
                    anomalies_found += 1

# Check table cells as well
for t_idx, table in enumerate(doc.tables):
    for r_idx, row in enumerate(table.rows):
        for c_idx, cell in enumerate(row.cells):
            clean_cell = " ".join(cell.text.split())
            matches = placeholder_pattern.findall(cell.text)
            if matches:
                print(f"[Placeholder] Table {t_idx} R{r_idx} C{c_idx}: matches={matches} -> {repr(clean_cell[:120])}")
                anomalies_found += 1
            for cat, terms in checks.items():
                for term in terms:
                    if cat == "ligatures":
                        if re.search(term, clean_cell, re.IGNORECASE):
                            print(f"[{cat.upper()}] Table {t_idx} R{r_idx} C{c_idx}: matched pattern {term} -> {repr(clean_cell[:120])}")
                            anomalies_found += 1
                    else:
                        if term.lower() in clean_cell.lower():
                            print(f"[{cat.upper()}] Table {t_idx} R{r_idx} C{c_idx}: matched '{term}' -> {repr(clean_cell[:120])}")
                            anomalies_found += 1

print(f"\nValidation finished. Total anomalies found: {anomalies_found}")
if anomalies_found > 0:
    sys.exit(1)
else:
    sys.exit(0)
