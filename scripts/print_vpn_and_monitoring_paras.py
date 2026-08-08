import docx
import sys

sys.stdout.reconfigure(encoding='utf-8')
doc = docx.Document("../Ugwulebo_Lesile_Ngozi_MIT_AWS_GCP_Final_Thesis.docx")

for idx in range(2512, 2963):
    if idx < len(doc.paragraphs):
        para = doc.paragraphs[idx]
        text = para.text.strip()
        if text.startswith("4.12.") or text.startswith("4.13.") or text.startswith("4.14.") or text.startswith("4.15.") or text.startswith("4.16"):
            print(f"P{idx}: {repr(para.text)}")
