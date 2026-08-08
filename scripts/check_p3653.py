import docx

d1 = docx.Document("../Ugwulebo_Lesile_Ngozi_MIT_AWS_GCP_Updated_Network(3).docx")
d2 = docx.Document("../Ugwulebo_Lesile_Ngozi_MIT_AWS_GCP_Final_Thesis.docx")

print("Original P3653:")
print(repr(d1.paragraphs[3653].text))

print("\nFinal P3653:")
print(repr(d2.paragraphs[3653].text))
