import docx
import sys

sys.stdout.reconfigure(encoding='utf-8')
doc = docx.Document("../Ugwulebo_Lesile_Ngozi_MIT_AWS_GCP_Final_Thesis.docx")

for idx in range(1380, 1410):
    if idx < len(doc.paragraphs):
        para = doc.paragraphs[idx]
        print(f"P{idx}: {repr(para.text)}")
