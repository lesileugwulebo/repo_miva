import docx

doc = docx.Document("../Ugwulebo_Lesile_Ngozi_MIT_AWS_GCP_Updated_Network(3).docx")
for idx in range(4047, 4060):
    if idx < len(doc.paragraphs):
        safe_text = doc.paragraphs[idx].text.encode('ascii', errors='replace').decode('ascii')
        print(f"P{idx}: {repr(safe_text)}")
