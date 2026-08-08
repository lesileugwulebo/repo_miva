import docx
import sys

sys.stdout.reconfigure(encoding='utf-8')
doc = docx.Document("../Ugwulebo_Lesile_Ngozi_MIT_AWS_GCP_Final_Thesis.docx")

for idx in range(1804, 1943):
    if idx < len(doc.paragraphs):
        para = doc.paragraphs[idx]
        text = para.text.strip()
        # Print headings, file headers, and start/end of code snippets
        if text.startswith("4.6") or text.startswith("File:") or "resource " in text or "direction " in text or "priority " in text or "name " in text or text == "}":
            print(f"P{idx}: {repr(para.text)}")
