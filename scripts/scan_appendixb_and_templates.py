import docx

doc_path = "../Ugwulebo_Lesile_Ngozi_MIT_AWS_GCP_Final_Thesis.docx"
doc = docx.Document(doc_path)

print(f"Total paragraphs: {len(doc.paragraphs)}")

print("\n--- Scanning for Template Paragraph Labels & Draft Strings ---")
for idx, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    if any(k in text.lower() for k in ["failover-analysis", "template", "will be inserted", "actual results are entered", "final results should"]):
        if not any(k in text.lower() for k in ["cloud platforms", "cloud adoption", "cloud formation"]):
            print(f"P{idx}: {text.encode('ascii', 'ignore').decode('ascii')}")

print("\n--- Scanning for Appendix B Content ---")
for idx, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    if "appendix b" in text.lower() or "evidence completion" in text.lower() or "checklist" in text.lower():
        print(f"P{idx}: {text.encode('ascii', 'ignore').decode('ascii')}")

print("\n--- Scanning for Chapter Capitalization Remnants ---")
for idx, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    if "Will Describe The Actual Implementation" in text or "Will Present The" in text or "Designed In This" in text:
        print(f"P{idx}: {text.encode('ascii', 'ignore').decode('ascii')}")
