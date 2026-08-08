import docx
import sys

sys.stdout.reconfigure(encoding='utf-8')
doc = docx.Document("../Ugwulebo_Lesile_Ngozi_MIT_AWS_GCP_Updated_Network(3).docx")

for idx in range(4510, 4555):
    if idx < len(doc.paragraphs):
        para = doc.paragraphs[idx]
        text = para.text
        if "depend" in text.lower() or "placeholders" in text.lower() or "evidence" in text.lower() or "completed" in text.lower():
            print(f"P{idx}: {repr(text)}")
