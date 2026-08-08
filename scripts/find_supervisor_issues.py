import docx

doc_path = "../Ugwulebo_Lesile_Ngozi_MIT_AWS_GCP_Final_Thesis.docx"
doc = docx.Document(doc_path)

print(f"Total paragraphs in document: {len(doc.paragraphs)}")

print("\n--- Scanning Abstract ---")
for idx in range(min(50, len(doc.paragraphs))):
    text = doc.paragraphs[idx].text
    if "abstract" in text.lower() or "evaluated" in text.lower() or "placeholder" in text.lower() or "fabricated" in text.lower():
        print(f"P{idx}: {text[:150]}")

print("\n--- Scanning Chapter 2 Note ---")
for idx in range(100, 300):
    text = doc.paragraphs[idx].text
    if "supervisor" in text.lower() or "revision prepared" in text.lower():
        print(f"P{idx}: {text}")

print("\n--- Scanning Chapter 6 Contradictions ---")
for idx in range(len(doc.paragraphs) - 500, len(doc.paragraphs)):
    text = doc.paragraphs[idx].text
    if any(k in text.lower() for k in ["deliberately uses", "outstanding element", "production-grade", "expressroute", "implement all four"]):
        print(f"P{idx}: {text}")

print("\n--- Scanning Terminology Issues ---")
for idx in range(len(doc.paragraphs)):
    text = doc.paragraphs[idx].text
    if "microsoft aws" in text.lower() or "expressroute" in text.lower() or "every field marked" in text.lower() or "all placeholders in chapter five" in text.lower():
        print(f"P{idx}: {text}")
