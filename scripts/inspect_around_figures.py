import docx

doc = docx.Document("../Ugwulebo_Lesile_Ngozi_MIT_AWS_GCP_Updated_Network(3).docx")

indices = [3798, 3828, 3892, 4010, 4054, 4117]

for idx in indices:
    print(f"\n=== Around P{idx} ===")
    for j in range(idx - 2, idx + 3):
        if 0 <= j < len(doc.paragraphs):
            print(f"P{j}: '{doc.paragraphs[j].text}'")
