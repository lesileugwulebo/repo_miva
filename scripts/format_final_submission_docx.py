import os
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

doc_path = "../Ugwulebo_Lesile_Ngozi_MIT_AWS_GCP_Final_Thesis.docx"
print(f"Loading document for academic formatting polish: {doc_path}")
doc = docx.Document(doc_path)

# Helper to set cell background color
def set_cell_background(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    tcPr.append(shd)

# Helper to set cell margins/padding
def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>')
    tcPr.append(tcMar)

print("Applying academic typography and paragraph formatting...")

# Set document margins (1 inch on all sides)
for section in doc.sections:
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

# Format paragraphs
for para in doc.paragraphs:
    text = para.text.strip()
    if not text and len(para.runs) == 0:
        continue

    style_name = para.style.name.lower()
    
    # Heading 1
    if style_name.startswith("heading 1") or text.startswith("Chapter "):
        para.paragraph_format.space_before = Pt(14)
        para.paragraph_format.space_after = Pt(8)
        para.paragraph_format.line_spacing = 1.15
        for run in para.runs:
            run.font.name = "Calibri"
            run.font.size = Pt(18)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A) # Dark Navy
            
    # Heading 2
    elif style_name.startswith("heading 2") or (text and text[0].isdigit() and "." in text[:5] and text.count(".") == 1):
        para.paragraph_format.space_before = Pt(12)
        para.paragraph_format.space_after = Pt(6)
        para.paragraph_format.line_spacing = 1.15
        for run in para.runs:
            run.font.name = "Calibri"
            run.font.size = Pt(14)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)
            
    # Heading 3
    elif style_name.startswith("heading 3") or (text and text[0].isdigit() and "." in text[:6] and text.count(".") >= 2):
        para.paragraph_format.space_before = Pt(10)
        para.paragraph_format.space_after = Pt(4)
        para.paragraph_format.line_spacing = 1.15
        for run in para.runs:
            run.font.name = "Calibri"
            run.font.size = Pt(12)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0x33, 0x41, 0x55)
            
    # Captions
    elif text.startswith("Figure ") or text.startswith("Table "):
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.paragraph_format.space_before = Pt(6)
        para.paragraph_format.space_after = Pt(4)
        for run in para.runs:
            run.font.name = "Calibri"
            run.font.size = Pt(10)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)
            
    # Figure Analysis Text
    elif para.runs and para.runs[0].italic and "Figure " in text:
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        para.paragraph_format.space_before = Pt(2)
        para.paragraph_format.space_after = Pt(10)
        for run in para.runs:
            run.font.name = "Calibri"
            run.font.size = Pt(10.5)
            run.font.color.rgb = RGBColor(0x47, 0x55, 0x69)
            
    # Normal Body Paragraphs
    else:
        if para.alignment == WD_ALIGN_PARAGRAPH.LEFT or para.alignment == WD_ALIGN_PARAGRAPH.JUSTIFY:
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(6)
        para.paragraph_format.line_spacing = 1.15
        for run in para.runs:
            if not run.font.name:
                run.font.name = "Calibri"
            if not run.font.size:
                run.font.size = Pt(11)

print("Formatting tables with executive academic styling...")
for table in doc.tables:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Format header row
    if len(table.rows) > 0:
        header_row = table.rows[0]
        trPr = header_row._tr.get_or_add_trPr()
        trPr.append(parse_xml(f'<w:tblHeader {nsdecls("w")}/>'))
        for cell in header_row.cells:
            set_cell_background(cell, "1E293B")
            set_cell_margins(cell, top=120, bottom=120, left=150, right=150)
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_before = Pt(2)
                p.paragraph_format.space_after = Pt(2)
                for run in p.runs:
                    run.font.name = "Calibri"
                    run.font.size = Pt(10)
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                    
    # Format data rows
    for r_idx, row in enumerate(table.rows[1:], start=1):
        bg_color = "F8FAFC" if r_idx % 2 == 1 else "FFFFFF"
        for cell in row.cells:
            set_cell_background(cell, bg_color)
            set_cell_margins(cell, top=80, bottom=80, left=120, right=120)
            for p in cell.paragraphs:
                p.paragraph_format.space_before = Pt(1)
                p.paragraph_format.space_after = Pt(1)
                for run in p.runs:
                    run.font.name = "Calibri"
                    run.font.size = Pt(9.5)

print(f"Saving formatted thesis document to: {doc_path}")
doc.save(doc_path)
print("Academic document formatting completed successfully!")
