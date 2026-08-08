import docx
import sys

sys.stdout.reconfigure(encoding='utf-8')
doc = docx.Document("../Ugwulebo_Lesile_Ngozi_MIT_AWS_GCP_Final_Thesis.docx")

print("Checking paragraphs in GCP sections...")
for idx in range(1694, 2260):
    if idx < len(doc.paragraphs):
        para = doc.paragraphs[idx]
        text = para.text.strip()
        # Print headings, file headers, and start/end of code snippets
        if text.startswith("4.") or text.startswith("File:") or idx in [1697, 1702, 1704, 1709, 1714, 1717, 1727, 1765, 1791, 1804, 1805, 1806, 1833, 1852, 1891, 1910, 1943, 1944, 2018, 2019, 2020, 2031, 2032, 2075, 2114, 2185, 2219, 2257]:
            print(f"P{idx}: {repr(para.text)}")
