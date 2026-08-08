import docx
import sys

sys.stdout.reconfigure(encoding='utf-8')
doc = docx.Document("../Ugwulebo_Lesile_Ngozi_MIT_AWS_GCP_Updated_Network(3).docx")

for idx in range(2930, 2963):
    if idx < len(doc.paragraphs):
        para = doc.paragraphs[idx]
        print(f"P{idx}: {repr(para.text)}")
