import docx
import re
import sys

sys.stdout.reconfigure(encoding='utf-8')

doc = docx.Document("../Ugwulebo_Lesile_Ngozi_MIT_AWS_GCP_Final_Thesis.docx")

search_terms = {
    "azure_resource": ["awsrm_", "azurerm_", "resource_group", "resource group", "Log Analytics", "Virtual WAN", "system-assigned", "Network VPC Firewall Rules"],
    "disclaimers": ["does not invent", "pending actual execution", "numerical result fields", "remain placeholders"],
    "classification": ["Highly Secure and Resilient Production-Grade", "Highly Secure and Resilient Production-Grade Multi-Cloud Topology"],
    "citation": ["Verdet et al. (2025)", "Verdet et al., 2025", "Verdet (2025)"],
    "ligatures": [r"\bdi er\b", r"\bo icial\b", r"\btra ic\b", r"\bo boarding\b", r"\bdi erent\b", r"\bdi erences\b"],
    "template_text": ["Chapter Five Will Present", "Chapter Two Reviewed", "Reviewed More Than", "Chapter Six Will Present"],
    "mixups": ["Security Groups in GCP", "NSG in AWS", "Security Group in GCP", "NSG in the AWS"],
    "duplicated_headings": ["TESTING, RESULTS, AND EVALUATION", "TESTING, RESULTS AND EVALUATION"]
}

print("Searching for issues in the document...")
for i, para in enumerate(doc.paragraphs):
    text = para.text
    clean_text = " ".join(text.split())
    
    # Check simple substrings
    for category, terms in search_terms.items():
        for term in terms:
            if category == "ligatures":
                if re.search(term, clean_text, re.IGNORECASE):
                    print(f"[{category}] P{i}: matched pattern {term} -> {repr(clean_text[:120])}")
            else:
                if term.lower() in clean_text.lower():
                    print(f"[{category}] P{i}: matched '{term}' -> {repr(clean_text[:120])}")

print("\nScanning table cells for issues...")
for t_idx, table in enumerate(doc.tables):
    for r_idx, row in enumerate(table.rows):
        for c_idx, cell in enumerate(row.cells):
            clean_cell = " ".join(cell.text.split())
            for category, terms in search_terms.items():
                for term in terms:
                    if category == "ligatures":
                        if re.search(term, clean_cell, re.IGNORECASE):
                            print(f"[{category}] Table {t_idx} R{r_idx} C{c_idx}: matched pattern {term} -> {repr(clean_cell[:120])}")
                    else:
                        if term.lower() in clean_cell.lower():
                            print(f"[{category}] Table {t_idx} R{r_idx} C{c_idx}: matched '{term}' -> {repr(clean_cell[:120])}")
