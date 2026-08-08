import docx
import sys

sys.stdout.reconfigure(encoding='utf-8')
doc = docx.Document("../Ugwulebo_Lesile_Ngozi_MIT_AWS_GCP_Final_Thesis.docx")

for idx in range(2842, 2963):
    if idx < len(doc.paragraphs):
        para = doc.paragraphs[idx]
        text = para.text.strip()
        if text.startswith("4.15.") or "resource " in text or "workspace" in text or "log_analytics" in text or "action_group" in text or "monitor_data" in text or "awsrm_" in text or text == "}":
            print(f"P{idx}: {repr(para.text)}")
