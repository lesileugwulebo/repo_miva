import docx
import sys

sys.stdout.reconfigure(encoding='utf-8')
doc = docx.Document("../Ugwulebo_Lesile_Ngozi_MIT_AWS_GCP_Updated_Network(3).docx")

for idx in range(4541, 5000):
    if idx < len(doc.paragraphs):
        para = doc.paragraphs[idx]
        text = para.text
        if "depend" in text.lower() or "placeholder" in text.lower() or "evidence" in text.lower() or "completed" in text.lower() or "outstanding" in text.lower():
            # Filter out non-disclaimers
            if any(k in text.lower() for k in ["depends on", "evidence is required", "outstanding element", "placeholder", "pending"]):
                print(f"P{idx}: {repr(text)}")
