import docx
import sys

sys.stdout.reconfigure(encoding='utf-8')
doc = docx.Document("../Ugwulebo_Lesile_Ngozi_MIT_AWS_GCP_Final_Thesis.docx")

current_file = None
start_idx = None

print("Scanning for code block paragraph ranges...")
for idx, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    if text.startswith("File:"):
        if current_file:
            print(f"File: {current_file} -> Range: P{start_idx} to P{idx-2}")
        current_file = text
        start_idx = idx + 1
    # Check if a new section starts which ends the code block
    elif text.startswith("4.") and current_file:
        print(f"File: {current_file} -> Range: P{start_idx} to P{idx-2}")
        current_file = None
        start_idx = None

if current_file:
    print(f"File: {current_file} -> Range: P{start_idx} to P{len(doc.paragraphs)-1}")
