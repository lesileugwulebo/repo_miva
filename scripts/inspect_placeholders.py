import docx
import re

doc = docx.Document("../Ugwulebo_Lesile_Ngozi_MIT_AWS_GCP_Updated_Network(3).docx")

print("Searching for bracketed placeholders...")
placeholders = []
pattern = re.compile(r"\[.*?\]")

# Search paragraphs
for i, para in enumerate(doc.paragraphs):
    matches = pattern.findall(para.text)
    for m in matches:
        placeholders.append((f"Paragraph {i}", m))

# Search tables
for t_idx, table in enumerate(doc.tables):
    for r_idx, row in enumerate(table.rows):
        for c_idx, cell in enumerate(row.cells):
            matches = pattern.findall(cell.text)
            for m in matches:
                placeholders.append((f"Table {t_idx} Row {r_idx} Col {c_idx}", m))

print(f"Found {len(placeholders)} placeholders:")
for loc, p in placeholders:
    print(f"- {loc}: {p}")
