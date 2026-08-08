import docx
import sys

sys.stdout.reconfigure(encoding='utf-8')
doc = docx.Document("../Ugwulebo_Lesile_Ngozi_MIT_AWS_GCP_Final_Thesis.docx")

figure_paras = [51, 53, 248, 374, 613, 858, 3470]

for idx in figure_paras:
    if idx < len(doc.paragraphs):
        para = doc.paragraphs[idx]
        print(f"P{idx}: text={repr(para.text)}")
        # Check if the paragraph has runs with drawings/inline_shapes
        has_run_image = False
        for run in para.runs:
            if "drawing" in run._element.xml:
                has_run_image = True
                break
        print(f"  Has drawing element in XML: {has_run_image}")
        # Also check paragraphs immediately after or before it
        for offset in [-1, 1]:
            n_idx = idx + offset
            if 0 <= n_idx < len(doc.paragraphs):
                n_para = doc.paragraphs[n_idx]
                has_img = False
                for run in n_para.runs:
                    if "drawing" in run._element.xml:
                        has_img = True
                if has_img:
                    print(f"  Adjacent P{n_idx} has drawing element!")
