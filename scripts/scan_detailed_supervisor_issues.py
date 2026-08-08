import docx

doc_path = "../Ugwulebo_Lesile_Ngozi_MIT_AWS_GCP_Final_Thesis.docx"
doc = docx.Document(doc_path)

print(f"Total paragraphs: {len(doc.paragraphs)}")

print("\n--- Scanning for 'Insert' instructions ---")
insert_count = 0
for idx, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    if "insert" in text.lower() and any(k in text.lower() for k in ["screenshot", "prowler", "scoutsuite", "terminal", "chart", "figure", "table"]):
        print(f"P{idx}: {text}")
        insert_count += 1
print(f"Total 'Insert' instructions found: {insert_count}")

print("\n--- Scanning for Azure syntax remnants in AWS blocks ---")
for idx, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    if any(k in text for k in ["VpnGw1AZ", "RouteBased", "active_active", "Microsoft. (2026). AWS Site-to-Site"]):
        print(f"P{idx}: {text}")

print("\n--- Scanning for Duplicate Headings & Odd Capitalization ---")
for idx, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    if text in ["SYSTEM IMPLEMENTATION", "TESTING, RESULTS, AND EVALUATION"] or "Will Present The" in text:
        print(f"P{idx}: {text}")

print("\n--- Scanning for Proposal / Future Tense ---")
will_count = 0
for idx, para in enumerate(doc.paragraphs):
    text = para.text
    if "will be used" in text or "will verify" in text or "will include" in text or "will pursue" in text:
        will_count += 1
print(f"Total future/proposal tense sentences found: {will_count}")
