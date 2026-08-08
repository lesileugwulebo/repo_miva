import docx
import re
import sys

# Ensure stdout is in UTF-8
sys.stdout.reconfigure(encoding='utf-8')

doc = docx.Document("../Ugwulebo_Lesile_Ngozi_MIT_AWS_GCP_Final_Thesis.docx")

print("Checking for remaining placeholders in final document...")
pattern = re.compile(r"\[[A-Za-z0-9_/ -]{3,}\]")
remaining = 0

for i, para in enumerate(doc.paragraphs):
    matches = pattern.findall(para.text)
    for m in matches:
        # Ignore normal code/metadata brackets
        if not any(k in para.text for k in ["[Unit]", "[Service]", "[Install]"]):
            # Use ascii representation or print safely
            safe_text = para.text[:60].encode('ascii', errors='replace').decode('ascii')
            print(f"- P{i}: {m} (Text: {safe_text}...)")
            remaining += 1

for t_idx, table in enumerate(doc.tables):
    for r_idx, row in enumerate(table.rows):
        for c_idx, cell in enumerate(row.cells):
            matches = pattern.findall(cell.text)
            for m in matches:
                print(f"- T{t_idx} R{r_idx} C{c_idx}: {m}")
                remaining += 1

print(f"Validation finished. Total remaining placeholders: {remaining}")
if remaining > 0:
    sys.exit(1)
