import docx
import re

doc = docx.Document("../Ugwulebo_Lesile_Ngozi_MIT_AWS_GCP_Updated_Network(3).docx")

with open("../results/placeholders_list.txt", "w", encoding="utf-8") as f:
    f.write("=== PARAGRAPH PLACEHOLDERS ===\n")
    for i, para in enumerate(doc.paragraphs):
        # Look for [ or ]
        if "[" in para.text or "]" in para.text:
            f.write(f"P{i}: {para.text}\n")
            
    f.write("\n=== TABLE CELL PLACEHOLDERS ===\n")
    for t_idx, table in enumerate(doc.tables):
        for r_idx, row in enumerate(table.rows):
            for c_idx, cell in enumerate(row.cells):
                if "[" in cell.text or "]" in cell.text:
                    f.write(f"T{t_idx} R{r_idx} C{c_idx}: {cell.text}\n")

print("Placeholders exported successfully to results/placeholders_list.txt")
