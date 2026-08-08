import docx

doc_path = "../Ugwulebo_Lesile_Ngozi_MIT_AWS_GCP_FINAL_SUBMISSION.docx"
doc = docx.Document(doc_path)

print(f"Validating final submission file: {doc_path}")
print(f"Total Paragraphs: {len(doc.paragraphs)}")
print(f"Total Tables: {len(doc.tables)}")

anomalies = 0
for idx, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    if "insert" in text.lower() and any(k in text.lower() for k in ["screenshot", "prowler", "scoutsuite", "terminal", "chart"]):
        print(f"Anomaly [INSTRUCTION]: P{idx}: {text}")
        anomalies += 1
    elif "placeholder" in text.lower() or "tbd" in text.lower() or "todo" in text.lower():
        print(f"Anomaly [PLACEHOLDER]: P{idx}: {text}")
        anomalies += 1

print(f"\nFinal Validation Finished. Total Anomalies Found: {anomalies}")
