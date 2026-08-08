import docx

doc_path = "../Ugwulebo_Lesile_Ngozi_MIT_AWS_GCP_FINAL_SUBMISSION_READY.docx"
doc = docx.Document(doc_path)

print("--- PARAGRAPHS P4550 TO P4610 ---")
for idx in range(4550, min(4615, len(doc.paragraphs))):
    print(f"P{idx}: {doc.paragraphs[idx].text}")
